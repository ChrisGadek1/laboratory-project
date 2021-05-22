from django.http import HttpResponse
from django.shortcuts import render
from django.core.files.storage import FileSystemStorage
from docx.oxml.ns import qn
from datetime import datetime
from docx.shared import Pt

from samples.forms import SampleForm, ResearchForm
import os
from samples.models import Sampling, WIJHARS, MetodAndNorm, ControlType, DeliveryWay, Type, Research, ResearchStatus
from docx import Document, table
from docx.oxml import OxmlElement
import copy
# Create your views here.


def choose_mode(request, *args, **kwargs):
    if request.user.is_authenticated:
        return render(request, 'reports/choose_mode.html', {})
    else:
        return render(request, 'main_not_logged.html', {})


def add_template(request, *args, **kwargs):
    context = {}
    if request.user.is_authenticated:
        fs = FileSystemStorage()
        if request.method == "POST" and request.POST.get("mode", "") == "add" and request.FILES['myfile']:
            myfile = request.FILES['myfile']
            filename = fs.save(myfile.name, myfile)
            uploaded_file_url = fs.url(filename)
        elif request.method == "POST"and request.POST.get("mode", "") == "delete":
            name = request.POST.get("choose-to-delete", "")
            fs.delete(name)

        samples_labels = {v: k for k, v in SampleForm.Meta.labels.items()}
        research_labels = {v: k for k, v in ResearchForm.Meta.labels.items()}
        context = {
            "samples_fields": samples_labels,
            "research_fields": research_labels,
            "reports": os.listdir("media/")
        }
        return render(request, 'reports/add_template.html', context)
    else:
        return render(request, 'main_not_logged.html', {})


def generate_report(request, *args, **kwargs):
    if request.user.is_authenticated:

        context = {}
        context["samples"] = Sampling.objects.all()
        context["reports"] = os.listdir("media/")
        if request.method == "POST":
            sample = Sampling.objects.get(code=request.POST.get("sample", ""))
            sample_values = {}
            for prop in SampleForm.Meta.labels.keys():
                value = ""
                if Sampling.objects.get(code=request.POST.get("sample", "")).__dict__.get(prop) is not None:
                    value = sample.__dict__.get(prop)
                else:
                    if prop == "WIJHARS":
                        value = WIJHARS.objects.get(id=sample.WIJHARS_id).name
                    elif prop == "control_type":
                        value = ControlType.objects.get(id=sample.control_type_id).name
                    elif prop == "sampling_method":
                        value = MetodAndNorm.objects.get(id=sample.sampling_method_id).name
                    elif prop == "sample_delivery":
                        value = DeliveryWay.objects.get(id=sample.sample_delivery_id).name
                    elif prop == "type":
                        value = Type.objects.get(id=sample.type_id).name
                if prop == "is_OK":
                    if value == "YES": value = "Tak"
                    else: value = "Nie"
                sample_values[prop] = str(value)
            document = Document("media/"+request.POST.get("template", ""))


            #inserting data about samples:
            for paragraph in document.paragraphs:
                for prop in sample_values.keys():
                    if "{{"+prop+"}}" in paragraph.text:
                        paragraph.text = paragraph.text.replace("{{"+prop+"}}", sample_values[prop])
            for table_v in document.tables:
                for column in table_v.columns:
                    for cell in column.cells:
                        for prop in sample_values.keys():
                            if "{{" + prop + "}}" in cell.text:
                                cell.text = cell.text.replace("{{" + prop + "}}", sample_values[prop])

            #inserting data about researches:

            researches = list(Research.objects.filter(sampling=sample).all())
            props = []
            if len(researches) > 0:
                props = researches[0].__dict__.keys()
            insert_flag = False
            result_type_sep = False
            numerical_result = False
            columns = []
            cells = []
            rows_to_delete = []
            merge_flag = True
            for table_v in document.tables:
                for row in table_v.rows:
                    if not insert_flag:
                        for cell in row.cells:
                            if "{{init_researches}}" in cell.text:
                                insert_flag = True
                                result_type_sep = False
                                rows_to_delete.append((table_v, row))
                            elif "{{init_researches_num}}" in cell.text:
                                insert_flag = True
                                result_type_sep = True
                                numerical_result = True
                                rows_to_delete.append((table_v, row))
                            elif "{{init_researches_char}}" in cell.text:
                                insert_flag = True
                                result_type_sep = True
                                numerical_result = False
                                rows_to_delete.append((table_v, row))
                    else:
                        cell_index = 0
                        for cell in row.cells:
                            cells.append(cell)
                            columns.append([])
                            has_prop_flag = False
                            for prop in props:
                                if "{{" + prop + "}}" in cell.text:
                                    columns[cell_index].append(prop)
                                    has_prop_flag = True
                            if not has_prop_flag:
                                if "{{incr}}" in cell.text:
                                    columns[cell_index].append("incr")
                                else:
                                    columns[cell_index].append(cell.text)
                            cell_index += 1
                        rows_to_delete.append((table_v, row))
                        #print(table_v._element.xml)
                        for it, research in enumerate(researches, start=1):
                            if result_type_sep and (numerical_result and not research.result.isnumeric()):
                                continue
                            elif result_type_sep and (not numerical_result and research.result.isnumeric()):
                                continue
                            row_tmp = table_v.add_row()
                            new_row = row_tmp.cells
                            cells_copy = copy.deepcopy(cells)
                            #word manipulation:
                            requirement_columns = {"law": 1, "specification": 2, "producent": 3}
                            requirement_number = 0
                            for i in range(len(columns)):
                                tc = new_row[i]._tc
                                tcPr = tc.tcPr
                                borders = OxmlElement('w:tcBorders')
                                border_list = ['w:left', 'w:right', 'w:top', 'w:bottom']
                                for border_element in border_list:
                                    element = OxmlElement(border_element)
                                    element.set(qn('w:val'), 'single')
                                    element.set(qn('w:sz'), '4')
                                    element.set(qn('w:space'), '0')
                                    element.set(qn('w:color'), '000000')
                                    borders.append(element)
                                tcPr.append(borders)
                                for j, replaceable in enumerate(columns[i]):
                                    if replaceable in props:
                                        to_replace = " "
                                        if replaceable == "requirements":
                                            requirement_number += 1
                                        if replaceable != "requirements" or requirement_columns[research.__dict__.get("requirementsType")] == requirement_number:
                                            to_replace = research.__dict__.get(replaceable)
                                        text = cells_copy[i].text
                                        if "{{" in text:
                                            new_row[i].text = text.replace("{{"+replaceable+"}}", str(to_replace))
                                        else:
                                            new_row[i].text = str(to_replace)
                                        cells_copy[i].text = new_row[i].text

                                    elif replaceable == "incr":
                                        new_row[i].text = str(it)
                                    else:
                                        new_row[i].text = replaceable

                        insert_flag = False
                        result_type_sep = False
                        numerical_result = False
                        columns = []
                        cells = []
                        break
            for data in rows_to_delete:
                if data[1]._element.getparent() is not None:
                    data[1]._element.getparent().remove(data[1]._element)
            document.save("generated_reports/report.docx")
            response = HttpResponse(content_type='application/vnd.openxmlformats-officedocument.wordprocessingml.document')
            response['Content-Disposition'] = 'attachment; filename=report.docx'
            document.save(response)
            return response

        return render(request, 'reports/reports.html', context)
    else:
        return render(request, 'main_not_logged.html', {})


def generate_time_report(request, *args, **kwargs):
    if request.user.is_authenticated:
        context = {}
        if request.method == "POST":
            document = Document("time_report/raport_czasowy.docx")
            begin_date = request.POST.get("input-date", "")
            end_date = request.POST.get("end-date", "")
            samples = Sampling.objects.all()
            researches = Research.objects.all()

            if begin_date == '':
                context = {"error_begin": "To pole nie może być puste"}
                return render(request, 'reports/generate_time_report.html', context)
            if end_date == '':
                context = {"error_end": "To pole nie może być puste"}
                return render(request, 'reports/generate_time_report.html', context)

            def filter_date_samples(date1):
                if date1.delivery_date is None:
                    return False
                else:
                    date1_obj = datetime.strptime(str(date1.delivery_date), "%Y-%m-%d")
                    end_date_obj = datetime.strptime(str(end_date), "%Y-%m-%d")
                    begin_date_obj = datetime.strptime(str(begin_date), "%Y-%m-%d")
                    return begin_date_obj <= date1_obj <= end_date_obj

            def filter_date_researches(date1):
                if date1.start_date is None:
                    return False
                else:
                    date1_obj = datetime.strptime(str(date1.start_date), "%Y-%m-%d")
                    end_date_obj = datetime.strptime(str(end_date), "%Y-%m-%d")
                    begin_date_obj = datetime.strptime(str(begin_date), "%Y-%m-%d")
                    return begin_date_obj <= date1_obj <= end_date_obj

            samples = list(filter(filter_date_samples, list(samples)))
            researches = list(filter(filter_date_researches, list(researches)))
            data_samples = {}
            for sample in samples:
                if data_samples.get(sample.type) is None:
                    data_samples[sample.type] = 1
                else:
                    data_samples[sample.type] += 1

            data_researches = {}
            for research in researches:
                if data_researches.get(research.status) is None:
                    data_researches[research.status] = 1
                else:
                    data_researches[research.status] += 1

            def modify_cell(tcl, text):
                tc = tcl._tc
                tcPr = tc.tcPr
                borders = OxmlElement('w:tcBorders')
                border_list = ['w:left', 'w:right', 'w:top', 'w:bottom']
                for border_element in border_list:
                    element = OxmlElement(border_element)
                    element.set(qn('w:val'), 'single')
                    element.set(qn('w:sz'), '4')
                    element.set(qn('w:space'), '0')
                    element.set(qn('w:color'), '000000')
                    borders.append(element)
                tcPr.append(borders)
                tcl.text = text

            p = document.paragraphs[0]
            p.text = str(p.text).replace("{{init_date}}", str(begin_date))
            p.text = str(p.text).replace("{{end_date}}", str(end_date))
            p.style.font.size = Pt(14)

            for i, table_v in enumerate(document.tables, start=1):
                if i == 1:
                    for type_s, number in data_samples.items():
                        row_tmp = table_v.add_row()
                        new_row = row_tmp.cells
                        modify_cell(new_row[0], str(type_s))
                        modify_cell(new_row[1], str(number))
                else:
                    for type_s, number in data_researches.items():
                        row_tmp = table_v.add_row()
                        new_row = row_tmp.cells
                        modify_cell(new_row[0], str(type_s))
                        modify_cell(new_row[1], str(number))

            document.save("time_report/report.docx")
            response = HttpResponse(content_type='application/vnd.openxmlformats-officedocument.wordprocessingml.document')
            response['Content-Disposition'] = 'attachment; filename=report.docx'
            document.save(response)
            return response
        return render(request, 'reports/generate_time_report.html', context)
    else:
        return render(request, 'main_not_logged.html', {})